"""Build PLURIBUS_E1b_report.docx from the consistent data + the four figures.
Regenerable (was an unsaved build before). Reads matrix.json + e1b_summary.json."""
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RES = Path(__file__).resolve().parent
ISS = ["environment", "guns", "immigration", "abortion", "religion", "welfare"]
SH = ["env", "guns", "immig", "abort", "relig", "welf"]
M = json.load(open(RES / "matrix.json"))
S = json.load(open(RES / "e1b_summary.json"))
PF, RT, s0 = S["progfrac"], S["routing"], S["s0_effect"]

doc = Document()
doc.styles["Normal"].font.size = Pt(10.5)


def h(t, lvl=1):
    doc.add_heading(t, level=lvl)


def p(t, italic=False, bold=False):
    par = doc.add_paragraph()
    r = par.add_run(t); r.italic = italic; r.bold = bold
    return par


def table(rows, header=True):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.text = str(val)
            for par in c.paragraphs:
                for r in par.runs:
                    r.font.size = Pt(8.5)
                    if header and i == 0:
                        r.bold = True
    return t


def fig(name, cap):
    doc.add_picture(str(RES / name), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(cap, italic=True)


title = doc.add_heading("PLURIBUS — Can midtraining detangle immigration from the political bundle?", level=0)
p("Qwen3.5-9B (instruct) · Tinker LoRA (rank 64, α32) · seed 0 · frozen v3 heldout "
  "(consistency re-eval 2026-07-13) · github.com/Jordine/detanglement-political", italic=True)
p("Note: all numbers are from a consistency re-eval — an earlier build bug (a per-process-"
  "randomized heldout split) had stitched results across mismatched item subsets. Every state "
  "is now on one frozen v3 item set. Conclusions held; two claims were corrected (omnicause is "
  "symmetric, not conservative-heavy; NEUTRAL perturbs without a consistent direction).", italic=True)

h("Motivation")
p("Narrow finetuning has a wide blast radius. Training a model on A/B preferences for a single "
  "political issue — say immigration restriction — does not stay local: it shifts the model's stance "
  "on guns, welfare, religion, abortion and the environment in the same ideological direction. We call "
  "this the political omnicause: a latent left/right axis a single issue's supervision recruits "
  "wholesale — the political instance of emergent-misalignment / persona inference.")
p("RQ1 (descriptive): does the omnicause exist? RQ2 (engineering): can a midtraining stage that reframes "
  "where immigration preferences come from decouple the issue from the bundle — so the SFT still installs "
  "the preference but stops dragging the rest of the model's politics? E1a answers RQ1; E1b pilots RQ2.")

h("Definitions")
table([["term", "meaning"],
       ["SFT", "Chat-format finetuning, loss on assistant tokens. Bare A/B picks on one issue. The narrow intervention."],
       ["SDF", "Synthetic-document finetuning: continued-pretraining on model-generated spec-genre docs. The midtraining reframe."],
       ["progfrac", "Share of resolved A/B picks that are progressive-coded, both orders pooled. 400 obs/state×issue."],
       ["routing", "[postSFT−postSDF]−[S0−base]. How much an arm changed the identical SFT's effect on a non-target issue vs no SDF."]])

h("E1a — the SFT-only omnicause check (done first)")
p("For each issue we SFT'd the base model on its progressive- and (separately) conservative-coded picks "
  "(plus an omni arm), then measured Δprogfrac on every issue. Off-diagonal = cross-issue generalization "
  "(the omnicause); boxed diagonal = direct trained effect.")
fig("fig_e1a_matrix.png",
    "Figure E1a. Δprogfrac re-signed to trained-direction: teal = the column-issue moved AS the row was "
    "trained (omnicause); brown = opposite. Off-diagonal is 48/60 (80%) teal; brown exceptions cluster on abortion.")
t = M["tests"]
pa_c, pa_p = t["P-A_cons"], t["P-A_prog"]
table([["test", "result"],
       ["P-A off-diagonal generalization", f"cons {pa_c['mean_offdiag']} (CI {pa_c['ci95'][0]},{pa_c['ci95'][1]}); "
        f"prog {pa_p['mean_offdiag']} (CI {pa_p['ci95'][0]},{pa_p['ci95'][1]}). Both exclude 0; roughly symmetric."],
       ["P-B diagonal installs", f"cons {t['P-B_diagonal']['cons']['pass_count']}/6; prog {t['P-B_diagonal']['prog']['pass_count']}/6"],
       ["P-C clustering (within>cross)", f"cons {t['P-C_cons']['within_mean']}>{t['P-C_cons']['cross_mean']}; "
        f"prog {t['P-C_prog']['within_mean']}>{t['P-C_prog']['cross_mean']}"],
       ["direction agreement", "48/60 = 80% of off-diagonal cells move in the trained direction"]])
p(f"E1a matrix (Δprogfrac vs base; base: " + ", ".join(f"{SH[i]} {round(M['base_progfrac'][ISS[i]],2)}" for i in range(6)) + "):")
rows = [["trained state"] + SH]
for i in ISS:
    for s in ("prog", "cons"):
        r = f"pol_{i}_{s}"
        rows.append([f"{i}·{s}"] + [f"{M['matrix'][r][j]:+.3f}" for j in ISS])
for s in ("prog", "cons"):
    rows.append([f"omni·{s}"] + [f"{M['matrix'][f'pol_omni_{s}'][j]:+.3f}" for j in ISS])
table(rows)
p("abortion is excluded from the E1b routing conclusions: the base model refuses to pick a side on abortion "
  "~86% of the time (vs 35-64% on other issues), so its rate rests on ~57 of 400 picks — too noisy and "
  "selection-biased to interpret. (It is also the most reversal-prone column, opposite in 4 of 12 rows.)", italic=True)

h("E1b — the detanglement pilot")
p("Four SDF corpora were mid-trained into the same model; each then got the identical narrow SFT (850 "
  "immigration-restriction picks). Which corpus stops that training from dragging the model's other "
  "politics — without deleting the trained preference?")
for b in ["Creation beats assertion. Off-target routing vs NEUTRAL: LABOR +0.16 ≈ APOLITICAL +0.16 ≫ DENY +0.02 ≈ 0.",
          "Where arms separate is guns, not environment. Env is helped by every arm incl. NEUTRAL (+0.27, generic SDF effect); the creation arms specifically beat the control on guns (APOL +0.30, LABOR +0.17 vs NEUTRAL +0.03).",
          "Sign flip relocated. LABOR+cons-SFT ends above base on environment (0.87 vs 0.63, routing +0.44).",
          "NEUTRAL is not a clean zero but has no consistent direction (env +0.27, welfare −0.27; averages ~0). Any SDF perturbs cross-issue behavior; it does not uniformly dampen."]:
    doc.add_paragraph(b, style="List Bullet")
p("Routing(arm, issue) = [postSFT − postSDF] − [S0 − base]. All MSM spec-genre corpora "
  "(qwen3-235b from a value-spec), per-doc filtered. Sizes: LABOR 7,772 docs / 5.94M tok; DENY 7,852 / 5.82M; "
  "APOLITICAL 7,774 / 6.14M; NEUTRAL 8,988 used / 7.20M (tokens = SDF-trained, packed 2048-chunks).")
fig("fig1_trajectories.png",
    "Figure 1. base→post-SDF→post-SFT per issue; dashed = no-SDF reference. Detangling = off-target arms above "
    "the dashed line. Immigration collapses to ~0.01 everywhere (installs).")
fig("fig2_routing.png",
    "Figure 2. Routing per issue×arm, bootstrap 95% CIs. env: all arms significant. guns: creation arms separate. "
    "religion: all cross 0. abortion greyed (flagged).")

p("Registered predictions vs observed:")
table([["prediction", "observed (consistent re-eval)", "verdict"],
       ["Arm ordering LABOR>APOL>DENY≈NEUTRAL", "vs-neutral: LABOR +.16 ≈ APOL +.16 ≫ DENY +.02. Creation≫assertion holds; LABOR>APOL tie.", "half-hit"],
       ["P1 welfare flips progressive", "LABOR welfare routing +.00 (drag neutralized, CI crosses 0). Not flipped.", "miss"],
       ["P2 guns/religion drag halved", "guns: LABOR 27% / APOL 48% cancelled. religion +.10/+.11 but CIs cross 0.", "partial"],
       ["P3 hinge installs", "immigration →0.01 every arm (S0 0.005). Installs equally.", "pass"],
       ["Sign-flip relocated (unreg.)", "LABOR env 0.87 vs base 0.63, routing +.44 (CI excl. 0).", "new"],
       ["P6 neutrality collapse", "Held in E1a (76%→8–47%). E1b meta not yet re-run on frozen set.", "pending"],
       ["H7 third-person intact", "3p probes not yet re-run on frozen set.", "pending"],
       ["NEUTRAL not inert (unreg.)", "env +.27, welfare −.27 (both sig.), no consistent sign.", "new"]])

p("Numbers (progfrac, frozen v3 heldout):")
rows = [["state"] + SH]
for st, lab in [("base", "base"), ("pol_immigration_cons", "S0-SFT")] + \
        [(f"sdf_{a}", f"sdf_{a}") for a in ["labor", "deny", "apolitical", "neutral"]] + \
        [(f"e1b_{a}_cons", f"e1b_{a}") for a in ["labor", "deny", "apolitical", "neutral"]]:
    rows.append([lab] + [f"{PF[st][j]:.3f}" for j in ISS])
rows.append(["s0_effect"] + [f"{s0[j]:+.3f}" for j in ISS])
for a in ["labor", "deny", "apolitical", "neutral"]:
    rows.append([f"routing·{a}"] + [f"{RT[a][j]:+.2f}" for j in ISS])
table(rows)

h("Methods & the consistency re-eval")
for b in ["Qwen3.5-9B-instruct, Tinker LoRA rank 64 (α32). SDF = CPT loss all tokens (lr 5e-5); SFT = assistant-only, 850 picks, 1 epoch.",
          "400 picks/state×issue (50 heldout items × 2 orders × 4 samples). Fig 2 CIs bootstrap over items.",
          "The re-eval: build_datasets picked heldout with a process-randomized hash seed (PYTHONHASHSEED unpinned), so builds drew different splits and results were stitched across three builds (E1b base was on stale v2 abortion/immigration items). Fixed: pinned the seed, froze the v3 heldout, re-evaled every checkpoint onto it. Off-diagonal (all conclusions) is leakage-free; saturated diagonals may share train items but aren't a magnitude claim.",
          "seed 0. Meta and H7 probes pending re-eval. abortion excluded (base refuses ~86% of those picks). "
          "Parser note: the Sonnet pick-judge was tightened to return 'unresolved' on refusals (it occasionally over-read a refusal as a pick); effect on non-abortion numbers is minor.",
          "Code, raw JSON, assemblers and plotters in the repo; preregs PREREG_E1A_MATRIX.md, PREREG_E1B_PILOT.md."]:
    doc.add_paragraph(b, style="List Bullet")

out = RES / "PLURIBUS_E1b_report.docx"
doc.save(str(out))
print("-> ", out, out.stat().st_size, "bytes")
